# -*- coding: utf-8 -*-

import os
import re
import sys
import zipfile

import xml.etree.ElementTree as ET

import convertUtil
import convertWord

from docx import Document
import docx

# The version using docx

# TIMESTAMP for version information.
TIMESTAMP = "Version 2018-06-28"

# Documentation here:
# https://docs.python.org/2/library/xml.etree.elementtree.html
# https://docs.python.org/2/library/zipfile.html

debug_output = True

# Flag controls if the conversion removes
# structure from replaced text when drawings are found
# in the .docx input.
removeOldText = False
# More aggressive removal of grandparents of empty text blocks.
removeOldTextParents = False

# If enabled, replaces old fonts in list in fontsTable.xml
convertFontsTable = True

# To keep track of data for an entire paragraph, to pass to conversion.
class paragraphData():
  def __init__(self):
    self.formatList = []
    self.textList = []
    self.textsize = 0
    self.startPoint = []

    self.paragraphText = ""
    self.paragraphTextStarts = []
    self.paragraphTextElements = []
    self.paragraphFonts = []

  def adddata(self, textNode, formatNode):
    self.formatList.append(formatNode)
    self.textList.append(textNode)
    self.startPoint.append(self.textsize)
    if textNode and textNode.text:
      self.textsize += len(textNode.text)

# Version 2. Copy old doc, then modify
class convertDocx():
  def __init__(self, input_path, output_dir, converter, debug=False):
    self.oldFonts = []
    self.input_path = input_path
    self.output_dir = output_dir
    self.converter = converter
    self.old_fonts = converter.oldFonts  # List of font names
    self.unicode_font = converter.unicodeFont
    self.debug = debug

    # True if paragraph text should be converted as a unit
    self.accumulate_text = False

    # If enabled, replaces old fonts in list in fontsTable.xml
    self.convertFontsTable = True

    # Pieces of the document
    self.paragraphs = []
    self.drawings = []
    self.text_boxes = []
    self.parent_map = {}
    self.paragraphs_converted = {}

    # Pieces for paragraphs
    self.paragraph_runs = []
    self.rprFormatData = []
    self.superscriptNode = None

    # Copy old doc into new path.
    self.document = Document(input_path)
    if self.output_dir:
      # String the directory tree to the file, substituting the output
      fileIn = os.path.split(self.input_path)[1]
      baseWOextension = os.path.splitext(fileIn)[0]
    else:
      baseWOextension = os.path.splitext(self.input_path)[0]
    self.outpath = os.path.join(self.output_dir, baseWOextension + '_unicode.docx')
    # Save the unchanged copy
    self.document.save(self.outpath)

  def processDocx(self):
    if self.debug:
      print('Convert2 processDocx path = %s, output_dir = %s\n' % (
        self.input_path, self.output_dir))
    # Try simple things.
    paragraphs = self.document.paragraphs
    for para in paragraphs:
      self.converter.processParagraphRuns(para)

    sections = self.document.sections
    print("Document has %s sections" % len(sections))
    for section in sections:
      header = section.header
      for para in header.paragraphs:
        self.converter.processParagraphRuns(para)
      footer = section.footer
      for para in header.paragraphs:
        self.converter.processParagraphRuns(para)

    tables = self.document.tables
    for table in tables:
      rows = table.rows
      for row in rows:
        for cell in row.cells:
          paragraphs = cell.paragraphs
          for para in paragraphs:
            self.converter.processParagraphRuns(para)

    self.document.save(self.outpath)
    return

  def tryFontUpdate(self, newzip, oldFontList, unicodeFont):
    filename = 'word/fontTable.xml'
    docXML = newzip.read(filename)  # A file-like object
    return self.parseFontTable(docXML, oldFontList)

  # Looks at text parts of the DOCX data, extracting each.
  def parseDocXML(self, docfile_name, docxml,
                  saveConversion=False, outpath=None, isString=False):
    if isString:
      tree = ET.fromstring(docxml)
      root = tree
    else:
      tree = ET.parse(self.input_path)
      root = tree.getroot()

    # Accumulate the paragraphs, drawing text, and text boxes
    drawingContentText = []
    textBoxText = []
    # Should we check for tables, too?
    for p in tree.iter():
      if re.search('}p$', p.tag):
        # Paragraphs in the whole document
        self.paragraphs.append(p)
      elif re.search('}drawing$', p.tag):
        pi = p.iter()
        for node in pi:
          if re.search('main}t$', node.tag):
            # Can I just convert this now?
            if node.text:
              drawingContentText.append(node)
      elif re.search('}txbxContent$', p.tag):
        pi = p.iter()
        for node in pi:
          if re.search('main}t$', node.tag):
            # Can I just convert this now?
            if node.text:
              textBoxText.append(node)

    print('@@@@@@ %d draw text' % len(drawingContentText))
    print('@@@@@@ %d text box text' % len(textBoxText))
    print('@@@@@@ %d paragraphs' % len(self.paragraphs))
    print('@@@@@@ docfile_name = %s' % docfile_name)

    if not drawingContentText and not textBoxText and docfile_name == 'word/document.xml':
      # Handle this with only DocX functions.
      doc = convertWord.convertOneDoc(self.input_path, self.converter)
      return doc

    # This contains drawings and possibly other things that can't be parsed with the docx
    # library. Process all the paragraphs and text items.
    # Compute the parent items for possible deletions.
    # http://elmpowered.skawaii.net/?p=74
    self.parent_map = dict((c, p) for p in tree.iter() for c in p)

    convertCount = 0

    allEmptiedTextElements, convertCount = self.processXmlParagraphs()

    # TODO: remove all emptied text elements.
    if self.accumulate_text:
      self.removeOldTextElements(allEmptiedTextElements)

    print(' %s: %d text items converted' % (docfile_name, convertCount))

    # Save the new document data.
    if isString:
      return ET.tostring(root, encoding='utf-8')

    if saveConversion:
      if not outpath:
        outpath = self.input_path
      tree.write(outpath)
    return tree

  # Process each of the paragraph nodes
  def processXmlParagraphs(self):
    allEmptiedTextElements = []
    paragraphCount = 0
    convertCount = 0
    self.paragraph_runs = []

    for para in self.paragraphs:
      newConvertedCount, emptiedElements, paragraph_text = self.handleXmlParagraph(para)

      convertCount += newConvertedCount
      allEmptiedTextElements.append(emptiedElements)
      print('@P@P@P Paragraph %d text = %s' % (paragraphCount, paragraph_text))
      paragraphCount += 1

    return allEmptiedTextElements, convertCount

  # Process the text of a single paragraph.
  def handleXmlParagraph(self, para):
    allEmptiedTextElements = []
    if para in self.paragraphs_converted:
      print('Paragraph already done: %s' % para)
      return 0, allEmptiedTextElements, ''
    paragraphText = ""
    paragraphTextStarts = []
    paragraphTextElements = []
    paragraphFonts = []
    conversions = 0
    # Accumulate data for all text in this paragraph.
    paragraph_info = paragraphData()

    textElements = []
    formatTextInfo = []
    self.rprFormatData = []
    collectedText = ''
    self.superscriptNode = None
    self.inEncodedFont = False
    actualFont = None
    fontNode = None
    fontFound = False
    convertedText = ''

    for pchild in para.iter('*'):
      if re.search('}r$', pchild.tag):
        # Look only at the rPr and <w:t> data
        for rchild in pchild.iter('*'):
          # Process <w:r>
          if re.search('}rPr', rchild.tag):
            self.processRtF(rchild, paragraph_info, self.rprFormatData)
            fontFound = False
            actualFont = None
            fontNode = None
            for rprchild in rchild.iter('*'):
              # Collect all the formatting information
              # TODO: revise returns.
              #fontFound, fontIndex = self.collectRprData(rprchild)
              # TODO: Move all to the above function

              # Process <w:rPr>
              # Gather formatting information as needed, e.g.,
              # superscript, subscript, style, underline, etc.
              if re.search('}vertAlign', rprchild.tag):
                superscriptNode = rprchild
              elif re.search('}rFonts', rprchild.tag):
                # Font info.
                fontFound = True
                fontNode = rprchild
                (newFontIndex, actualFont) = self.isOldFontNode(rprchild)
                if newFontIndex >= 0:
                  # In font encoded node
                  self.inEncodedFont = True
                  fontIndex = newFontIndex
                else:
                  # Check if we are switching out. If so, handle accumulated text
                  if self.inEncodedFont:
                    # TODO: Even if no text, reset the encoded font
                    if collectedText:
                      (newConvertedCount, emptiedElements) = \
                        self.processCollectedText(
                          collectedText,
                          textElements,
                          superscriptNode,
                          formatTextInfo,
                          fontIndex)
                      conversions += newConvertedCount
                      allEmptiedTextElements.append(emptiedElements)
                    collectedText = ''
                    self.rprFormatData = []
                    formatTextInfo = []
                    textElements = []
                    self.inEncodedFont = False
          elif re.search('}t', rchild.tag):
            # For all the paragraph text.
            paragraphTextStarts.append(len(paragraphText))
            self.paragraph_runs.append(rchild)
            if rchild.text:
              paragraphText += rchild.text
            paragraphTextElements.append(rchild)
            paragraphFonts.append(actualFont)
            paragraph_info.adddata(rchild, fontNode)

            try:
              treat_as_no_break = self.converter.checkContentsForMerge(rchild.text)
            except:
              treat_as_no_break = False

            if fontFound and (self.inEncodedFont or treat_as_no_break) and rchild.text:
              # Consider the text across the runs as a unit
              if self.accumulate_text:
                # Process <w:t>
                formatTextInfo.append([rchild.text, self.rprFormatData])
                collectedText += rchild.text
                self.rprFormatData = []
                textElements.append(rchild)
              else:
                # Convert the text of each run independently.
                rchild.text = self.converter.convertText(
                  rchild.text, fontTextInfo=None, fontIndex=0)
                convertedText += rchild.text
                conversions += 1
            else:
              notEncoded = rchild.text
              if notEncoded and False and debug_output:
                print(' &&& fontFound = %s, self.inEncodedFont = %s, actual = %s' %
                      (fontFound, self.inEncodedFont, actualFont))
                print('notEncoded = >%s<' % notEncoded)

    # End of the paragraph
    if paragraphText and self.accumulate_text:
      convertedText = self.converter.convertText(paragraphText, fontTextInfo=None,
                                                 fontIndex=0)
      if debug_output:
        print('     Starts = %s' % paragraphTextStarts)
        print('  --> Converted paragraph = %s' % convertedText)

    # TODO: Use the collect paragraph stuff to do better conversion, esp. sentence
    # and punctuation across font differences.
    # Use self.paragraph_runs and paragraphText to do sentence-level stuff
    # if needed.
    try:
      self.converter.processXmlParagraphRuns(self.paragraph_runs)
    except:
      pass

    self.paragraphs_converted[para] = True
    return conversions, allEmptiedTextElements, convertedText

  def removeOldTextElements(self, allElementsToRemove):
    count = 0
    if not removeOldText:
      return count

    # This may be causing corruption in the MS Word file structure.
    for group in reversed(allElementsToRemove):
      for item in reversed(group):
        parent = self.parent_map[item]
        parent.remove(item)

        if not removeOldTextParents:
          continue
        # Can I remove the parent of this, too?
        grandparent = self.parent_map[parent]
        if grandparent is not None:
          grandparent.remove(parent)
        count += 1

    # And probably remove the siblings and the empty parent, too.
    return count


  def fixElementAndParent(self, textElement, parent, newText):
    removeList = []
    oldText = textElement.text
    for item in parent.findall('*'):
      for child in item.findall('*'):
        if re.search('}rFonts', child.tag):
          attrib = child.attrib
          for key in attrib:
            if attrib[key] in self.old_fonts:
              attrib[key] = self.unicode_font
        elif re.search('}vertAlign', child.tag):
          # TODO: Fix! This is specific for Old Osage
          if (oldText == u'H' or oldText == u'\uf048'):
            keys = child.attrib.keys()
            if re.search('}val', keys[0]):
              child.attrib[keys[0]] = 'baseline'

    textElement.text = newText


  # Replace the text in the first text element with the converted
  # unicode string, remove text from other text elements in that,
  # batch, and remove the empty elements.
  # Should I reset the font in this function, too?
  def processCollectedText(self, collectedText, textElementList,
                           formatTextInfo, fontIndex):
    # type: (object, object, object, object, object, object) -> object
    clearedTextElements = []
    global debug_output

    # First, change the text
    if False and debug_output:
      print('** COLLECTED %s to Unicode. ' % collectedText)

    convertedText = self.converter.convertText(collectedText, fontTextInfo=formatTextInfo,
                                               fontIndex=fontIndex)
    # TODO: Add case conversion, if desired.

    convertedCount = 0
    if convertedText != collectedText:
      convertedCount = 1
      #if debug_output:
      #  print('  ++++ Converted: %s' % convertedText)
    else:
      if debug_output:
        print('  ---- Not converted: %s' % collectedText)

    # 1. Reset text in first element
    if not textElementList:
      print('!!!!!!!!!!!!!!! NO TEXT ELEMENT LIST')
      return 0

    try:
      parent = self.parent_map[textElementList[0]]

      # Fix font and superscripting
      self.fixElementAndParent(textElementList[0], parent, convertedText)
    except:
      print('Error with parent indexing')

    if self.superscriptNode:
      # This is special case for Osage. TODO: Move to Osage self.converter
      self.superscriptNode.val = 'baseline'

    # 2. Clear text in other elements
    for element in textElementList[1:]:
      element.text = ''
      clearedTextElements.append(element)

    # TODO: Consider removing cleared elements.
    return convertedCount, clearedTextElements

  def processRtF(self, rchild, paragraph_info, rprFormatData):
    # Deal with the font and text data for the paragraph.
    # Keep track of all the data for each chunk and
    self.inEncodedFont = False
    if re.search('}rPr', rchild.tag):
      fontFound = False
      actualFont = None
      fontNode = None
      for rprchild in rchild.iter('*'):
        # Collect all the formatting information
        self.rprFormatData.append(rprchild)

        # Process <w:rPr>
        if re.search('}vertAlign', rprchild.tag):
          self.superscriptNode = rprchild
        elif re.search('}rFonts', rprchild.tag):
          self.processFont(rprchild)        # Font info.
          fontFound = True
          (newFontIndex, actualFont) = self.isOldFontNode(rprchild)
          if newFontIndex >= 0:
            # In font encoded node
            self.inEncodedFont = True
            fontIndex = newFontIndex
          else:
            # Check if we are switching out. If so, handle accumulated text
            if self.inEncodedFont:
              # TODO: Even if no text, reset the encoded font
              if collectedText:
                (newConvertedCount, emptiedElements) = \
                  self.processCollectedText(collectedText,
                                            textElements,
                                            parent_map,
                                            formatTextInfo, fontIndex)
                convertCount += newConvertedCount
                allEmptiedTextElements.append(emptiedElements)
              collectedText = ''
              self.rprFormatData = []
              formatTextInfo = []
              textElements = []
              self.inEncodedFont = False
    elif re.search('}t', rchild.tag):
      processText(rchild, paragraphTextStarts, paragraphTextElements, paragraphFonts,
                  collectedText, self.rprFormatData, textElements,
                  fontFound, actualFont)

  def collectRprData(self, rprchild):
    self.rprFormatData.append(rprchild)

    fontFound = True
    fontIndex = -1
    fontNode = None

    # Process <w:rPr>
    # Gather formatting information as needed, e.g.,
    # superscript, subscript, style, underline, etc.
    if re.search('}vertAlign', rprchild.tag):
      self.superscriptNode = rprchild
    elif re.search('}rFonts', rprchild.tag):
      # Font info.
      fontFound = True
      fontNode = rprchild
      (newFontIndex, actualFont) = self.isOldFontNode(rprchild)
      if newFontIndex >= 0:
        # In font encoded node
        self.inEncodedFont = True
        fontIndex = newFontIndex
      else:
        # Check if we are switching out. If so, handle accumulated text
        if self.inEncodedFont:
          # TODO: Even if no text, reset the encoded font
          if collectedText:
            (newConvertedCount, emptiedElements) = \
              self.processCollectedText(
                collectedText,
                textElements,
                self.superscriptNode,
                formatTextInfo,
                fontIndex)
            conversions += newConvertedCount
            allEmptiedTextElements.append(emptiedElements)
          collectedText = ''
          self.rprFormatData = []
          formatTextInfo = []
          textElements = []
          self.inEncodedFont = False
    return fontFound, fontIndex

  def processFont(self, rprchild):
    fontFound = True
    (newFontIndex, actualFont) = self.isOldFontNode(rprchild)
    self.inEncodedFont = False

    # Check if we are switching out. If so, handle accumulated text
    if self.inEncodedFont:
      # TODO: Even if no text, reset the encoded font
      if collectedText:
        (newConvertedCount, emptiedElements) = \
          self.processCollectedText(collectedText,
                                    textElements,
                                    self.converter,
                                    formatTextInfo, newFontIndex)
        convertCount += newConvertedCount
        allEmptiedTextElements.append(emptiedElements)


  def processText(self, rchild, paragraphTextStarts, paragraphTextElements, paragraphFonts,
                  collectedText, rprFormatData, textElements,
                  fontFound, actualFont):
    paragraphTextStarts.append(len(paragraphText))
    if rchild.text:
      paragraphText += rchild.text
    paragraphTextElements.append(rchild)
    paragraphFonts.append(actualFont)

    treat_as_no_break = self.converter.checkContentsForMerge(rchild.text)
    if treat_as_no_break:
      # Prepare to handle this, too.
      print('Contents treated as no break: %s' % rchild.text)

    if fontFound and (self.inEncodedFont or treat_as_no_break) and rchild.text:
      # Process <w:t>
      v.append([rchild.text, rprFormatData])
      collectedText += rchild.text
      self.rprFormatData = []
      textElements.append(rchild)

    else:
      # Mostly a check for
      notEncoded = rchild.text
      if notEncoded and False and debug_output:
        print(' &&& fontFound = %s, self.inEncodedFont = %s, actual = %s' %
              (fontFound, self.inEncodedFont, actualFont))
        print('notEncoded = >%s<' % notEncoded)


  def isOldFontNode(self, node):
    # Look for "rFonts", and check if any font contains one of the old fonts
    # Returns index of font if found, else -1
    if re.search('}rFonts', node.tag):
      for key in node.attrib:
        fontIndex = 0
        for oldFont in self.old_fonts:
          if re.search(oldFont, node.attrib[key]):
            return (fontIndex, node.attrib[key])
          fontIndex += 1
    return (-1, node.attrib[key])


  def parseFontTable(self, docXML, oldFontList):
    tree = ET.fromstring(docXML)

    for node in tree.iter('*'):
      if re.search('}font$', node.tag):
        keys = node.attrib.keys()
        for key in keys:
          if re.search('}name', key) and node.attrib[key] in self.oldFonts:
            print( 'Replacing font %s with %s' % (node.attrib[key], self.unicode_font))
            node.attrib[key] = self.unicode_font
    return ET.tostring(tree, encoding='utf-8')


# For standalone and testing.
def main(argv):
  global debug_output

  args = convertUtil.parseArgs()

  paths_to_doc = args.filenames
  print('ARGS = %s' % args)

  for path in paths_to_doc:
    extension = os.path.splitext(path)[-1]
    if extension == '.docx':
      docxProcessor = convertDoc.convertDocx(path, args.output_dir, args.font, debug_output)
      docxProcessor.processDocx()
    else:
      print('!!! Not processing file %s !' % path)


if __name__ == "__main__":
  main(sys.argv)