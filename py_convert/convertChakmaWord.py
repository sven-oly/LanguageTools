# -*- coding: utf-8 -*-
#!/usr/bin/env python

from __future__ import absolute_import, division, print_function

import os
import re
import sys

# Read and process MS Word documents, converting old Chakma encoding into
# Unicode characters.

# https://openpyxl.readthedocs.io/en/default/tutorial.html

LANGUAGE = 'Chakma'

from docx import Document

import chakmaConversion

# Flag for handling all characters in an Old font.
convertAllInOldFontRange = True

debugFlag = False

def replFunc(matchObj):
  if matchObj.group(0):
    if notOsageLatinLower.search(matchObj.group(0)):
      return matchObj.group(0)
    else:
      return osageConversion.oldOsageToUnicode(matchObj.group(0))

# It uses the font to set the right conversion
def checkAndConvertText(textIn, fontName):

  if textIn[0] == '=':
    # Ignore function calls
    return textIn

  result = chakmaConversion.oldEncodingToUnicode(textIn, fontName)

  return result


def convertDoc(doc, unicodeFont, debugInfo=None):
  sections = doc.sections
  print ('  %d sections' % len(sections))

  paragraphs = doc.paragraphs
  print ('  %d paragraphs' % len(doc.paragraphs))

  if debugInfo:
    print ('  %d tables' % len(doc.tables))
    if doc.inline_shapes:
      print ('  %d inline_shapes' % len(doc.inline_shapes))
    if doc.part:
      print (' %s part' % dir(doc.part))

    print ('    doc dir: %s' % dir(doc))
    for section in sections:
      print ('Section = %s' % section)

  numConverts = 0
  notConverted = 0
  paraNum = 0
  for para in paragraphs:

    para_format = para.paragraph_format
    para_style = para.style
    para_alignment = para.alignment
    para_part = para.part

    if debugInfo:
      print ('  Paragraph %d' % paraNum)
      print ('    para format = %s' % para_format)
      print ('    para style = %s' % para_style)
      print ('    para alignment = %s' % para_alignment)
      print ('    para part = %s' % para_part)
      if para_part:
        print ('    inline_shapes = %s' % para_part.inline_shapes)

    runs = para.runs
    if debugInfo:
      print ('    %d runs in paragraph' % (len(runs)))
      print ('    paragraph text = %s' % (para.text))
    runNum = 0
    runNum = 1
    for run in runs:
      if len(run.text):
        if debugInfo:
          print ('  Run %d text(%d) =  >%s<' % (runNum, len(run.text), run.text))
        thisText = run.text
        fontObj = run.font
        fontName = fontObj.name

        if fontName not in chakmaConversion.FONTS_TO_CONVERT:
          print ('not found fontName = %s' % fontName)
          continue
        print ('Font found = %s' % fontName)
        if thisText:
          convertedText = checkAndConvertText(thisText, fontName)
          if thisText != convertedText:
            numConverts += 1
            run.text = convertedText
            fontObj.name = unicodeFont
          else:
            notConverted += 1
      runNum += 1
    paraNum += 1

  print ('  %d values converted to Unicode' % numConverts)
  return (numConverts, notConverted)


# Process one DOCX, substituting the
def convertOneDoc(path_to_doc, unicodeFont='NotoSansChakma-Regular',
                  outpath=None, isString=False):

  print ('Converting  in file: %s' % path_to_doc)

  doc = Document(path_to_doc)

  (numConverts, numNoteConverted) = convertDoc(doc, unicodeFont, debugInfo=debugFlag)

  if numConverts:
    newName = os.path.splitext(path_to_doc)[0]
    unicode_path_to_doc = newName + '.unicode.docx'
    doc.save(unicode_path_to_doc)
    print ('  ** Saved new version to file %s\n' % unicode_path_to_doc)
  else:
    print ('  @@@ No conversion done, so no new file created.\n')


def processArgs(argv):
  if len(sys.argv) <= 1:
    print ('Usage:')
    print ('  convertWord.py inputFile.docx')
    print ('  convertWord.py inputFile1.docx inputFile2.docx ... ')
    print ('  convertWord.py -i fileWithFileNames')
    return None

  path_to_docs = []

  if len(argv) == 2:
    path_to_docs.append(sys.argv[1])
  else:
    if len(argv) == 3 and argv[1] == '-f':
      # Get the file containing conversion list and get all items.
      path_to_docs = convertUtil.infileToList(argv[2])
      if not path_to_docs:
        print ('Error: no contents found in file %s' %
               argv[2])
        return
    else:
      # Expect a list of files in the
      path_to_docs = [path for path in argv[1:]]

  return path_to_docs


def main(argv):

  if len(sys.argv) > 1:
    path_to_doc = sys.argv[1]
  else:
    print ('An input .docx file is required.')

  doc_list = processArgs(argv)

  convertFileCount = 0
  for doc_path in doc_list:
    convertOneDoc(doc_path)
    convertFileCount += 1
  print ('%d processed' % convertFileCount)


if __name__ == "__main__":
  main(sys.argv)
