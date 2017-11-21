# -*- coding: utf-8 -*-

from __future__ import absolute_import, division, print_function

import os
import re
import sys

# Read and process MS Word documents, converting Old Osage encoding into
# Unicode characters.

# https://openpyxl.readthedocs.io/en/default/tutorial.html

OfficialOsageFont = 'Official Osage Language'
FONTS_TO_CONVERT = [OfficialOsageFont]

from docx import Document

import convertUtil
import osageConversion

# Flag for handling all characters in an Old font.
convertAllInOldFontRange = True

# Rule for detecting Latin text or Old Osage font.
# Some Old Osage text is in Latin CAPS, but with lower case a, e, and o.
latinOsagePattern2 = r'[\^A-Z\[\]][A-Zaeo\[\]\^\\\'\/\._`,!]+'

# This identifies traditional Osage private use characters
traditionalOsageCharacters = ur'([\uf040-\uf05d]+)'

# To avoid converting English words
notOsageLatinLower = re.compile(r'[b-df-np-z]')

osageConvertPattern = latinOsagePattern2 + '|' + traditionalOsageCharacters

def replFunc(matchObj):
  if matchObj.group(0):
    if notOsageLatinLower.search(matchObj.group(0)):
      return matchObj.group(0)
    else:
      return osageConversion.oldOsageToUnicode(matchObj.group(0))

# Check for Osage text and convert the Osage parts of the strings.
# It assumes that the font has been detected.
def checkAndConvertText(textIn):

  if textIn[0] == '=':
    # Ignore function calls
    return textIn
  # if notOsageLatinLower.search(textIn):
  #  return textIn

  # Handle Latin and TraditionalOsage private use characters.
  if convertAllInOldFontRange:
    result = osageConversion.oldOsageToUnicode(textIn)
    return result
  else:
    # Attempt to get the Old Osage parts of the string.
    tryResult = re.subn(osageConvertPattern, replFunc, textIn)
    if tryResult[1] >= 1:
      return tryResult[0]
    else:
      return textIn


def convertDoc(doc, unicodeFont, debugInfo=None,
               extractedFile=None):
  sections = doc.sections
  print ('  %d sections' % len(sections))

  paragraphs = doc.paragraphs
  print ('  %d paragraphs' % len(doc.paragraphs))

  print ('  %d tables' % len(doc.tables))
  if debugInfo and doc.inline_shapes:
    print ('  %d inline_shapes' % len(doc.inline_shapes))
  if debugInfo and doc.part:
   print (' %s part' % dir(doc.part))

  if debugInfo:
    print ('    doc dir: %s' % dir(doc))
    for section in sections:
      print ('Section = %s' % section)

  numConverts = 0
  notConverted = 0
  paraNum = 0

  # Output line for extracted text and Python conversion
  extractedLine = []
  extractedCount = 0

  for para in paragraphs:

    para_format = para.paragraph_format
    para_style = para.style
    para_alignment = para.alignment
    para_part = para.part

    if debugInfo:
      print ('  Paragraph %d' % paraNum)
      print ('    para format = %s' % para_format)
      print ('    para style = %s' % para_style)
      print ('    para alignment = %s' % para_alignment)
      print ('    para part = %s' % para_part)
      if para_part:
        print ('    inline_shapes = %s' % para_part.inline_shapes)

    runs = para.runs
    print ('    %d runs in paragraph' % (len(runs)))
    print ('    paragraph text = %s' % (para.text))
    runNum = 0
    runNum = 1
    for run in runs:
      if len(run.text):
        if debugInfo:
          print ('  Run %d text(%d) =  >%s<' % (runNum, len(run.text), run.text))
        thisText = run.text
        fontObj = run.font
        fontName = fontObj.name
        if fontName not in FONTS_TO_CONVERT:
          continue
        if thisText:
          print ('run font = %s' % run.font.name)
          convertedText = checkAndConvertText(thisText)
          if thisText != convertedText:
            numConverts += 1
            run.text = convertedText
            fontObj.name = unicodeFont
            if extractedFile:
              extractedFile.write('%s\t%s\t%s\n' % (
                  fontName, thisText, convertedText))
          else:
            notConverted += 1
      runNum += 1
    paraNum += 1

  if extractedFile:
     extractedFile.close()

  print ('  %d values converted to Unicode' % numConverts)
  return (numConverts, notConverted)


# Process one DOCX, substituting the
def convertOneDoc(path_to_doc, unicodeFont='Pawhuska',
                  outpath=None, isString=False):

  print ('Converting Osage in file: %s' % path_to_doc)

  doc = Document(path_to_doc)

  if extractedFileName:
    extractedFile = codecs.open(extractedFileName, 'w', 'utf-8')
  else:
    extractedFile = None

  (numConverts, numNoteConverted) = convertDoc(doc, unicodeFont,
                                               extractedFile =extractedFile)

  if numConverts:
    newName = os.path.splitext(path_to_doc)[0]
    unicode_path_to_doc = newName + '.unicode.docs'
    doc.save(unicode_path_to_doc)
    print ('  ** Saved new version to file %s\n' % unicode_path_to_doc)
  else:
    print ('  @@@ No conversion done, so no new file created.\n')


def processArgs(argv):
  if len(sys.argv) <= 1:
    print ('Usage:')
    print ('  convertWordOsage.py inputFile.docx')
    print ('  convertWordOsage.py inputFile1.docx inputFile2.docx ... ')
    print ('  convertWordOsage.py -i fileWithFileNames')
    return None

  path_to_docs = []

  if len(argv) == 2:
    path_to_docs.append(sys.argv[1])
  else:
    if len(argv) == 3 and argv[1] == '-f':
      # Get the file containing conversion list and get all items.
      path_to_docs = convertUtil.infileToList(argv[2])
      if not path_to_docs:
        print ('Error: no contents found in file %s' %
               argv[2])
        return
    else:
      # Expect a list of files in the
      path_to_docs = [path for path in argv[1:]]

  return path_to_docs


def main(argv):

  if len(sys.argv) > 1:
    path_to_doc = sys.argv[1]
  else:
    path_to_doc = 'Harry Red Eagle Track 5.doc'

  doc_list = processArgs(argv)

  convertFileCount = 0
  for doc_path in doc_list:
    convertOneDoc(doc_path)
    convertFileCount += 1
  print ('%d files were processed' % convertFileCount)


if __name__ == "__main__":
  main(sys.argv)
