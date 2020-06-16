# -*- coding: utf-8 -*-
#!/usr/bin/env python3

from __future__ import absolute_import, division, print_function

import os
import re
import sys

# Read and process MS Word documents, converting old Cherokee encoding into
# Unicode characters.

# https://openpyxl.readthedocs.io/en/default/tutorial.html
from docx import Document

import adlamConversion
import convertUtil

import inspect

FONTS_TO_CONVERT = ['Fulfulde - Pulaar', 'MT Extra', 'Cherokee OLD',
                    'Cherokee;Cherokee2',
                    'Cherokee2']
# Flag for handling all characters in an Old font.
convertAllInOldFontRange = True

debugFlag = True  # False

# Set to True to get lower case conversion
toLowerCase = True  # False


# Check for text and convert the old font encoded parts of the strings.
# It assumes that the font has been detected.
def checkAndConvertText(textIn, converter):
  if not textIn or textIn[0] == '=':
    # Ignore function calls
    return textIn
  fontIndex = 0
  fontTextInfo = None
  result = converter.convertText(textIn, fontIndex=fontIndex)
  return result


class paragraphData():
  def __init__(self):
    self.runs = []
    self.formatList = []
    self.textList = []
    self.textsize = 0
    self.startPoint = []

  def adddata(self, run, textNode, formatNode):
    self.runs.append(run)
    self.formatList.append(formatNode)
    self.textList.append(textNode)
    self.startPoint.append(self.textsize)
    self.textsize += len(textNode.text)


class runInfo():
  def __init__(self, run):
    self.run = run
    self.sentence_start_offsets = []
    self.exclamation_mark_offsets = []
    self.question_mark_offsets = []
    self.stop_mark_offsets = []
    self.text = ''

# Creates a new document with converted text, adding in pieces as neede
# and flexibly handling text runs.
class copyConverter():
  def __init__(self, path_in, path_out=None):
    self.path_in = path_in
    self.doc_in = Document(self.path_in)

    styles = self.doc_in.styles
    print('styles (%d)= %s' % (len(styles), styles))
    for style in styles:
      print('  style = %s' % (style))

    if not path_out:
      newName = os.path.splitext(self.path_in)[0]
      path_out = newName + '.u2.docx'
    self.path_out = path_out

    self.converter = None
    # TODO FIX THIS
    self.unicode_font_out = 'Noto Sans Adlan New'

  def deleteParagraph(self, paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

  def convertParagraph(selfself, p):
    # TODO Fill this in.
    return

  def convertDoc(self, unicodeFont, debugInfo=None):
    # Creating a new document
    newdoc = Document()

    # TODO: set up style with new Unicode font
    for style in self.doc_in.styles:
      newdoc.styles.add_style(style.name, style.type)
      new_paragraph = newdoc.add_paragraph()
      new_paragraph.style = style
      print('Added style %s' % style.name)
      self.deleteParagraph(new_paragraph)

    for old_section in self.doc_in.sections:
      new_section = newdoc.add_section(start_type=old_section.start_type)

      new_section.page_height = old_section.page_height
      new_section.page_width = old_section.page_width
      new_section.orientation = old_section.orientation
      new_section.top_margin = old_section.top_margin
      new_section.gutter = old_section.gutter
      new_section.header_distance = old_section.header_distance
      new_section.footer_distance = old_section.footer_distance
      new_section.left_margin = old_section.left_margin
      new_section.right_margin = old_section.right_margin

      hparagraph = new_section.header.paragraphs[0]
      hparagraph.text = old_section.header.paragraphs[0].text

    p_num = 0
    for old_p in self.doc_in.paragraphs:
      if not old_p.text:
        print('P %d is empty' % p_num)

      p_num += 1

      # Add simple text paragraph, converted as a whole
      # newText = self.converter.convertText(old_p.text) if self.converter else old_p.text

      new_p = newdoc.add_paragraph()
      # Copy properties of the paragraph
      new_p.style = old_p.style
      old_p_format = old_p.paragraph_format
      new_p_format = new_p.paragraph_format

      # new_p.paragraph_format = old_p.paragraph_format
      new_p_format.line_spacing_rule = old_p_format.line_spacing_rule
      new_p_format.alignment = old_p_format.alignment
      new_p_format.left_indent = old_p_format.left_indent
      new_p_format.right_indent = old_p_format.right_indent
      new_p_format.first_line_indent = old_p_format.first_line_indent

      # Now get the runs of text, with compatibilities
      print('RUNS %2d old %2d new', (len(old_p.runs), len(new_p.runs)))
      new_runs = self.extractRuns(old_p, new_p)

      # Convert each run
      for run in new_p.runs:
        new_text = self.converter.convertText(run.text)
        run.text = new_text
        run.font.name = self.unicode_font_out
        run.rtl = self.converter.isRtl()

      # Now consider special punctuation and capitalization
      # TODO

      # For testing after combining and converting.
      self.analyzeRuns(new_runs)
      # TODO: process capitalization and punctuation adjustments.

    newdoc.save(self.path_out)

  def compareRunFonts(self, f1, f2):
    font_diffs = []
    if f1.all_caps != f2.all_caps:
      font_diffs.append('all_caps')
    if f1.bold != f2.bold:
      font_diffs.append('bold')
    if f1.bold != f2.bold:
      font_diffs.append('bold')
    # if f1.theme_color != f2.theme_color:
    #  font_diffs.append('theme_color')
    if f1.complex_script != f2.complex_script:
      font_diffs.append('complex_script')
    if f1.cs_bold != f2.cs_bold:
      font_diffs.append('cs_bold')
    if f1.cs_italic != f2.cs_italic:
      font_diffs.append('cs_italic')
    if f1.double_strike != f2.double_strike:
      font_diffs.append('double_strike')
    if f1.emboss != f2.emboss:
      font_diffs.append('emboss')
    if f1.hidden != f2.hidden:
      font_diffs.append('hidden')
    if f1.highlight_color != f2.highlight_color:
      font_diffs.append('highlight_color')
    if f1.imprint != f2.imprint:
      font_diffs.append('imprint')
    if f1.italic != f2.italic:
      font_diffs.append('italic')
    if f1.math != f2.math:
      font_diffs.append('math')
    if f1.name != f2.name:
      font_diffs.append('name')
    # if f1.theme_color != f2.theme_color:
    #  font_diffs.append('theme_color')
    if f1.outline != f2.outline:
      font_diffs.append('outline')
    if f1.rtl != f2.rtl:
      font_diffs.append('rtl')
    if f1.shadow != f2.shadow:
      font_diffs.append('shadow')
    if f1.size != f2.size:
      font_diffs.append('size')
    if f1.small_caps != f2.small_caps:
      font_diffs.append('small_caps')
    if f1.snap_to_grid != f2.snap_to_grid:
      font_diffs.append('snap_to_grid')
    if f1.spec_vanish != f2.spec_vanish:
      font_diffs.append('spec_vanish')
    if f1.strike != f2.strike:
      font_diffs.append('strike')
    if f1.subscript != f2.subscript:
      font_diffs.append('subscript')
    if f1.superscript != f2.superscript:
      font_diffs.append('superscript')
    if f1.underline != f2.underline:
      font_diffs.append('underline')
    if f1.web_hidden != f2.web_hidden:
      font_diffs.append('web_hidden')
    return font_diffs

  def runsCompatible(self, prev, current):
    # Check the characteristics for compatibility
    # This includes most attributes, but not rtl or font details
    # The checks may include looking at the contents of the run's text
    # compared with the output expected.
    if not prev:
      return False

      # Also check font properties such as size, color, style, etc.
    mergeable_text = self.converter.checkContentsForMerge(current.text)
    if mergeable_text:
      return True
    diffs = self.compareRunFonts(prev.font, current.font)
    # Here's where we check for important differences
    # print('Font diffs = %s' % (diffs))
    if not diffs:
      return True
    return False  # This will just copy runs.

  def copyFontData(self, new_run, old_run):
    old_font = old_run.font

    new_run.font.all_caps = old_font.all_caps
    new_run.font.bold = old_font.bold
    new_run.font.color.rgb = old_font.color.rgb
    new_run.font.color.theme_color = old_font.color.theme_color
    new_run.font.complex_script = old_font.complex_script
    new_run.font.cs_bold = old_font.cs_bold
    new_run.font.cs_italic = old_font.cs_italic
    new_run.font.double_strike = old_font.double_strike
    new_run.font.emboss = old_font.emboss
    new_run.font.hidden = old_font.hidden
    new_run.font.highlight_color = old_font.highlight_color
    new_run.font.imprint = old_font.imprint
    new_run.font.italic = old_font.italic
    new_run.font.math = old_font.math
    new_run.font.name = old_font.name
    new_run.font.no_proof = old_font.no_proof
    new_run.font.outline = old_font.outline
    new_run.font.rtl = old_font.rtl
    new_run.font.shadow = old_font.shadow
    new_run.font.size = old_font.size
    new_run.font.small_caps = old_font.small_caps
    new_run.font.snap_to_grid = old_font.snap_to_grid
    new_run.font.spec_vanish = old_font.spec_vanish
    new_run.font.strike = old_font.strike
    new_run.font.subscript = old_font.subscript
    new_run.font.superscript = old_font.superscript
    new_run.font.underline = old_font.underline
    new_run.font.web_hidden = old_font.web_hidden

  def extractRuns(self, old_p, new_p):
    # Get run data, combining runs as needed.
    p_data = paragraphData()
    runs = []
    prev_run = None
    empty_run_count = 0

    # TEMPORARY FOR TESTING
    for run in old_p.runs:
      # Compare attributes
      # If compatible with previous, append text

      if prev_run and self.runsCompatible(prev_run, run):
        print('Combining runs %s with %s' % (prev_run.text, run.text))
        prev_run.text += run.text
      else:
        if not run.text:
          empty_run_count += 1
        else:
          # else create new run with the style
          new_run = new_p.add_run(run.text)
          new_run.style = run.style
          new_run.bold = run.bold
          new_run.italic = run.italic
          new_run.underline = run.underline

          self.copyFontData(new_run, run)
          # print('Run %s text size = %s' % (len(runs), len(new_run.text)))
          runs.append(new_run)
          prev_run = new_run
    print('%d runs found. %d empty found' % (len(runs), empty_run_count))

    return runs

  def analyzeRuns(self, run_list):
    # Look for beginning of sentences for capitalization and insertion
    # of start of sentence marks.
    analyses = []

    end_sentence_pattern = r'([\.\?\!\u061F])($|\u0020)'

    # Find if there are ends of sentences with exclamation or question marks
    print('---------------')
    starting = True
    for run in run_list:
      analysis = runInfo(run)
      analysis.text = run.text
      sentence_parts = re.split(end_sentence_pattern, run.text)
      analyses.append(analysis)
      if starting:
        analysis.sentence_start_offsets.append(0)  # Actual first non-white-space
        starting = False
      match_iter = re.finditer(end_sentence_pattern, run.text)
      for match in match_iter:
        span = match.span()
        first = span[0]
        matched_char = match.groups(0)[0][0]
        # print('   match %s %s [%s]' % (match.span(), match.groups(), matched_char))
        if matched_char == '!':
          analysis.exclamation_mark_offsets.append(first)
        elif matched_char == '.':
          analysis.stop_mark_offsets.append(first)
        else:
          analysis.question_mark_offsets.append(first)
        if span[1] >= len(run.text) and (span[1] - span[0] > 1):
          starting = True  # The previous was

    # The last run is assumed to end a sentence.
    # TODO: Complete
    runnum = 0
    # for analysis in analyses:
    #   print(' Analysis %2d %s %s %s %s %s' % (runnum,
    #                                           analysis.sentence_start_offsets,
    #                                           analysis.stop_mark_offsets,
    #                                           analysis.exclamation_mark_offsets,
    #                                           analysis.question_mark_offsets,
    #                                           analysis.text,
    #                                           ))
    runnum += 1


# Handles sentence conversion if needed.
def convertParagraph(para, converter, unicodeFont, debugInfo=False):
  numConverts = 0
  notConverted = 0
  runs = para.runs
  if debugInfo:
    print('    %d runs in paragraph' % (len(runs)))
    print('    paragraph text = %s' % (para.text))
  runNum = 0
  runNum = 1
  fontsInRun = []
  for run in runs:
    if debugInfo:
      print('    run %d text = (%d) %s' % (runNum, len(run.text), run.text))
      print('run element = %s' % run._element)
      print('run parent = %s' % run._element.getparent())
    fontsInRun.append(run.font)

    if len(run.text):
      thisText = run.text
      fontObj = run.font
      fontName = fontObj.name
      # print('$$$$ name  = %s' % inspect.getmembers(fontObj.name))q
      if debugInfo:
        print('  Run #%1d in font >%s<. Text(%d) =  >%s<' % (
          runNum, fontName, len(run.text), run.text))

      if converter.forceFont:
        fontObj.name = unicodeFont

      if thisText:
        convertedText = checkAndConvertText(thisText, converter)
        if thisText != convertedText:
          numConverts += 1
          run.text = convertedText
          fontObj.name = unicodeFont
        else:
          notConverted += 1
    if debugInfo:
      print('Fonts in run %s' % fontsInRun)
    if len(fontsInRun) > 1:
      compareFonts(fontsInRun[0], fontsInRun[1])
    runNum += 1

    # Additional processing as needed.
  converter.processParagraphRuns(para)

  return numConverts


def convertTables(doc, converter, unicodeFont):
  # Look at the paragraphs in each cell of each table
  numConverted = 0

  for table in doc.tables:
    # Look at the cells
    row_num = 0
    for row in table.rows:
      for cell in table.row_cells(row_num):
        p = cell.paragraphs
        for p in cell.paragraphs:
          numConverted += convertParagraph(p, converter, unicodeFont)
      row_num += 1
  return numConverted


def convertDoc(doc, converter, unicodeFont, debugInfo=None):
  sections = doc.sections
  print('  %d sections' % len(sections))
  print('  %d tables' % len(sections))
  print('  unicodeFont = %s' % unicodeFont)
  paragraphs = doc.paragraphs
  print('  %d paragraphs' % len(doc.paragraphs))

  numConverts = 0
  notConverted = 0

  # Add font for Unicode
  for style in doc.styles:
    if style.name.find('Default') >= 0:
      print('DEFAULT %s' % style.font.name)
      style.font.name = unicodeFont
    print('Style %s: %s' % (style.name, style.type))

  # Headers and footers
  for section in sections:
    header = section.header
    for p in header.paragraphs:
      numConverts += convertParagraph(p, converter, unicodeFont)

  numConverts += convertTables(doc, converter, unicodeFont)

  paraNum = 0
  for para in paragraphs:
    num_converted = convertParagraph(para, converter, unicodeFont)
    numConverts += num_converted
    paraNum += 1

  print('  %d values converted to Unicode' % numConverts)
  return (numConverts, notConverted)


def compareFonts(f1, f2):
  # TODO: Figure this out.
  f1dir = inspect.getmembers(f1)
  f2dir = inspect.getmembers(f2)
  sameFont = True
  return sameFont


# Process one DOCX, substituting text in the old font with converted values.
def convertOneDoc(path_to_doc, converter, unicodeFont='Noto Sans Adlam New',
                  outpath=None, isString=False):
  print('Converting in file: %s' % path_to_doc)

  doc = Document(path_to_doc)

  (numConverts, numNoteConverted) = convertDoc(
    doc, converter, unicodeFont, debugInfo=debugFlag)

  if numConverts:
    newName = os.path.splitext(path_to_doc)[0]
    unicode_path_to_doc = newName + '.unicode.docx'
    doc.save(unicode_path_to_doc)
    print('  ** Saved new version to file %s\n' % unicode_path_to_doc)
  else:
    print('  @@@ No conversion done, so no new file created.\n')


def processArgs(argv):
  if len(sys.argv) <= 1:
    print('Usage:')
    print('  convertWord.py inputFile.docx')
    print('  convertWord.py inputFile1.docx inputFile2.docx ... ')
    print('  convertWord.py -i fileWithFileNames')
    return None

  path_to_docs = []

  if len(argv) == 2:
    path_to_docs.append(sys.argv[1])
  else:
    if len(argv) == 3 and argv[1] == '-f':
      # Get the file containing conversion list and get all items.
      path_to_docs = convertUtil.infileToList(argv[2])
      if not path_to_docs:
        print('Error: no contents found in file %s' %
              argv[2])
        return
    else:
      # Expect a list of files in the
      path_to_docs = [path for path in argv[1:]]

  return path_to_docs


def main(argv):
  if len(sys.argv) > 1:
    path_to_doc = sys.argv[1]
  else:
    print('An input .docx file is required.')

  doc_list = processArgs(argv)

  adlamFonts = [['Fulfulde - Aissata', 'arab'], ]
  convertFileCount = 0
  adlam_converter = adlamConversion.converter(adlamFonts)
  adlam_converter.setLowerMode(toLowerCase)
  adlam_converter.debug = False

  create_new_doc = False
  # Creates a copy
  for doc_path in doc_list:
    if create_new_doc:
      copy_converter = copyConverter(doc_path)
      copy_converter.converter = adlam_converter
      copy_converter.convertDoc('Noto Sans Adlam New')
    else:
      # Make modified copy.
      convertOneDoc(doc_path, adlam_converter)
    convertFileCount += 1

  print('%d processed' % convertFileCount)


if __name__ == "__main__":
  main(sys.argv)
