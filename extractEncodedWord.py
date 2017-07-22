# -*- coding: utf-8 -*-

from __future__ import absolute_import, division, print_function

import codecs
import os
import re
import sys

# Read and process MS Word documents, converting font encoding into
# Unicode characters.

# https://openpyxl.readthedocs.io/en/default/tutorial.html

OldFont = ''
FONTS_TO_CONVERT = {
    'Arjyaban Normal': {'encoding': 'Arjyaban', 'replace': 'RibengUni'},
    'Arjyaban CN': {'encoding': 'Arjyaban', 'replace': 'RibengUni'},
    'RajgirCN': {'encoding': 'Arjyaban', 'replace': 'RibengUni'},
    'Chakma(SuJoyan)': {'encoding': 'Sujoyan', 'replace': 'RibengUni'},
    'SutonnyMJ': {'encoding': 'Sutonny'},
    }

from docx import Document

import convertUtil
import conversion  # The converter(s) to use.

# Flag for handling all characters in an Old font.
convertAllInOldFontRange = True

debugFlag = False


# Check for font-encoded text and convert it to Unicode.
# It assumes that the font has been detected.
def checkAndConvertText(textIn, encoding):

  # Handle use characters.
  result = conversion.convertToUnicode(textIn, encoding)
  return result


def convertDoc(doc, unicodeFont, debugInfo=None, extractedFileName='extracted.tsv'):
  fonts_found = {}
  allMissingChars = {}
  sections = doc.sections
  print ('  %d sections' % len(sections))

  paragraphs = doc.paragraphs
  print ('  %d paragraphs' % len(doc.paragraphs))

  if debugInfo:
    print ('  %d tables' % len(doc.tables))
    if doc.inline_shapes:
      print ('  %d inline_shapes' % len(doc.inline_shapes))
    if doc.part:
      print (' %s part' % dir(doc.part))

    print ('    doc dir: %s' % dir(doc))
    for section in sections:
      print ('Section = %s' % section)

  numConverts = 0
  notConverted = 0
  paraNum = 0

  # Output line for extracted text and Python conversion
  extractedLine = []
  extractedFile = codecs.open(extractedFileName, 'w', 'utf-8')
  extractedCount = 0

  for para in paragraphs:

    para_format = para.paragraph_format
    para_style = para.style
    para_alignment = para.alignment
    para_part = para.part

    if debugInfo:
      print ('  Paragraph %d' % paraNum)
      print ('    para format = %s' % para_format)
      print ('    para style = %s' % para_style)
      print ('    para alignment = %s' % para_alignment)
      print ('    para part = %s' % para_part)
      if para_part:
        print ('    inline_shapes = %s' % para_part.inline_shapes)

    runs = para.runs
    if debugInfo:
      print ('    %d runs in paragraph' % (len(runs)))
      print ('    paragraph text = %s' % (para.text))
    runNum = 0
    runNum = 1

    # Should I replace the encoding find in the paragraph, too?
    for run in runs:
      fontObj = run.font
      fontName = fontObj.name
      if fontName in FONTS_TO_CONVERT:
        print ('Run font = %s. Text has %d chars' % (fontName, len(run.text)))
      if len(run.text):
        if debugInfo:
          print ('  Run %d text(%d) =  >%s<' % (runNum, len(run.text), run.text))
        thisText = run.text
        fontObj = run.font
        fontName = fontObj.name

        if fontName is None:
          # ??? Get the paragraph's style
          fontName = para.style.font.name
          fontObj = run.font
          print ('*** Getting font %s from Paragraph ***' % (fontName))

        if fontName not in fonts_found:
          # Find the characters in each font.
          fonts_found[fontName] = []
          if debugInfo:
            print ('FONT FOUND = %s' % fontName)
            print ('  text = %s' % thisText.encode('utf-8'))
        for t in thisText:
          if t not in fonts_found[fontName]:
            fonts_found[fontName].append(t)

        if fontName not in FONTS_TO_CONVERT:
          continue
        if thisText:
          # Record code points found in this encoding.

          encoding = FONTS_TO_CONVERT[fontName]['encoding']
          # print ('Encoding %s for %s' % (encoding, thisText.encode('utf-8')))
          (convertedText, missingChars) = conversion.convertToUnicode(thisText, encoding)

          if thisText != convertedText:
            numConverts += 1
            try:
              extractedFile.write('%s\t%s\t%s\n' % (
                  fontName, thisText, convertedText))

              run.text = convertedText
            except Exception as err:
              print ('Error = %s' % err)
              print ("Error in setting run.text with %s" % convertedText.encode('utf-8'))
          else:
            if debugInfo:
              print ('******** Not converted encoding = %s, text = >%s<' % (
                  encoding, thisText.encode('utf-8')))
            notConverted += 1

          if missingChars:
            for t in missingChars:
              if t in allMissingChars:
                allMissingChars[t] += 1
              else:
                allMissingChars[t] = 1
        # Change the font, even if there was no text changed.
        if fontName in FONTS_TO_CONVERT and 'replace' in FONTS_TO_CONVERT[fontName]:
          try:
            print ('  Reset font %s to %s ' % (fontObj.name, FONTS_TO_CONVERT[fontName]['replace']))
            fontObj.name = FONTS_TO_CONVERT[fontName]['replace']
          except Exception as err:
            print ('Problem replacing font in fontObj %s with %s' % (fontObj, fontName))

      runNum += 1
    paraNum += 1

  extractedFile.close()

  print ('  %d values converted to Unicode' % numConverts)
  for font in fonts_found:
    print ('font = %s' % font)
    print ('  chars = %s' % sorted(fonts_found[font]))

  print ('All Missing Chars = %s' % allMissingChars)
  return (numConverts, notConverted)


# Process one DOCX, substituting the
def convertOneDoc(path_to_doc, unicodeFont='RibengUni',
                  outpath=None, isString=False):

  print ('Converting text in file: %s' % path_to_doc)

  doc = Document(path_to_doc)

  newName = os.path.splitext(path_to_doc)[0]
  extractName = newName + '_extracted.tsv'
  (numConverts, numNotConverted) = convertDoc(
      doc, unicodeFont, debugInfo=debugFlag,
      extractedFileName=extractName)

  fonts_found = {}
  allMissingChars = {}
  sections = doc.sections
  print ('  %d sections' % len(sections))

  paragraphs = doc.paragraphs
  print ('  %d paragraphs' % len(doc.paragraphs))

  if debugFlag:
    print ('  %d tables' % len(doc.tables))
    if doc.inline_shapes:
      print ('  %d inline_shapes' % len(doc.inline_shapes))
    if doc.part:
      print (' %s part' % dir(doc.part))

    print ('    doc dir: %s' % dir(doc))
    for section in sections:
      print ('Section = %s' % section)

  numConverts = 0
  notConverted = 0
  paraNum = 0

  if numConverts:
    unicode_path_to_doc = newName + '.unicode.docx'
    doc.save(unicode_path_to_doc)
    print ('  ** Saved new version to file %s\n' % unicode_path_to_doc)
    print ('    uncoverted = %d' % numNotConverted)
  else:
    print ('  @@@ No conversion done, so no new file created.\n')


def processArgs(argv):
  if len(sys.argv) <= 1:
    print ('Usage:')
    print ('  convertWordChakma.py inputFile.docx')
    print ('  convertWordChakma.py inputFile1.docx inputFile2.docx ... ')
    print ('  convertWordChakma.py -i fileWithFileNames')
    return None

  path_to_docs = []

  if len(argv) == 2:
    path_to_docs.append(sys.argv[1])
  else:
    if len(argv) == 3 and argv[1] == '-f':
      # Get the file containing conversion list and get all items.
      path_to_docs = convertUtil.infileToList(argv[2])
      if not path_to_docs:
        print ('Error: no contents found in file %s' %
               argv[2])
        return
    else:
      # Expect a list of files in the
      path_to_docs = [path for path in argv[1:]]

  return path_to_docs


def main(argv):

  print ('ARGS = %s' % argv)
  if len(argv) > 1:
    path_to_doc = sys.argv[1]
  else:
    print ('Please provide an input file of type .docx')
    return

  doc_list = processArgs(argv)

  convertFileCount = 0
  for doc_path in doc_list:
    convertOneDoc(doc_path)
    convertFileCount += 1
  print ('%d processed' % convertFileCount)


if __name__ == "__main__":
  main(sys.argv)
